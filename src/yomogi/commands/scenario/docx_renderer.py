import re
from copy import deepcopy
from dataclasses import dataclass
from html import unescape
from importlib.resources import as_file, files

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

from yomogi.commands.scenario.models import (
    Conversation,
    DocumentTitle,
    Node,
    PartTitle,
    Scene,
    StageDirection,
    TextBlock,
)

FONT_SIZE_PT = 10.5
HEADING_FONT_SIZE_PT = 11.0
MINCHO_FONT = "游明朝"
GOTHIC_FONT = "游ゴシック"
DEFAULT_TEMPLATE = "scenario_classic.docx"
RED = RGBColor(0xFF, 0x00, 0x00)
BLUE = RGBColor(0x00, 0x00, 0xFF)


@dataclass(frozen=True)
class TextRun:
    text: str
    color: RGBColor | None = None


def _plain_text(text: object) -> str:
    without_tags = re.sub(r"<[^>]+>", "", str(text))
    return unescape(without_tags)


def _text_runs(text: object) -> list[TextRun]:
    source = str(text)
    runs: list[TextRun] = []
    position = 0
    pattern = re.compile(
        r'<span class="(?P<class>red|blue)">(?P<text>.*?)</span>'
    )
    for match in pattern.finditer(source):
        if match.start() > position:
            runs.append(TextRun(_plain_text(source[position : match.start()])))
        color = RED if match.group("class") == "red" else BLUE
        runs.append(TextRun(unescape(match.group("text")), color=color))
        position = match.end()

    if position < len(source):
        runs.append(TextRun(_plain_text(source[position:])))

    return [run for run in runs if run.text]


def _set_japanese_font(
    run,
    *,
    bold: bool = False,
    size: int = FONT_SIZE_PT,
    font_family: str = MINCHO_FONT,
) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_family
    fonts = run._element.rPr.rFonts
    for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attribute), font_family)
    fonts.set(qn("w:hint"), "eastAsia")


def _clear_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)


def _paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def _add_paragraph(
    paragraph,
    text: object,
    *,
    bold: bool = False,
    size: int = FONT_SIZE_PT,
    font_family: str = MINCHO_FONT,
    indent_chars: int = 0,
    hanging_chars: int = 0,
) -> None:
    _clear_paragraph(paragraph)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if indent_chars:
        paragraph.paragraph_format.left_indent = Pt(FONT_SIZE_PT * indent_chars)
    if hanging_chars:
        paragraph.paragraph_format.first_line_indent = Pt(
            -FONT_SIZE_PT * hanging_chars
        )
    for text_run in _text_runs(text):
        run = paragraph.add_run(text_run.text)
        _set_japanese_font(
            run,
            bold=bold,
            size=size,
            font_family=font_family,
        )
        if text_run.color:
            run.font.color.rgb = text_run.color


def _add_text_runs(
    paragraph,
    text: object,
    *,
    bold: bool = False,
    size: int = FONT_SIZE_PT,
    font_family: str = MINCHO_FONT,
) -> None:
    for text_run in _text_runs(text):
        run = paragraph.add_run(text_run.text)
        _set_japanese_font(
            run,
            bold=bold,
            size=size,
            font_family=font_family,
        )
        if text_run.color:
            run.font.color.rgb = text_run.color


def _add_conversation_paragraph(paragraph, block: Conversation) -> None:
    _clear_paragraph(paragraph)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(FONT_SIZE_PT)
    paragraph.paragraph_format.first_line_indent = Pt(-FONT_SIZE_PT)

    _add_text_runs(paragraph, block.speaker, font_family=GOTHIC_FONT)
    _add_text_runs(
        paragraph,
        f"{block.opening_mark}{block.text}{block.closing_mark}",
        font_family=MINCHO_FONT,
    )


def _add_blank_paragraph(paragraph) -> None:
    _clear_paragraph(paragraph)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def _load_template() -> Document:
    resource = files("yomogi.commands.scenario").joinpath(
        "templates",
        DEFAULT_TEMPLATE,
    )
    with as_file(resource) as template:
        document = Document(template)
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), MINCHO_FONT)
    return document


def render_docx(nodes: list[Node]) -> Document:
    document = _load_template()
    template_paragraph = document.paragraphs[0]
    template_element = deepcopy(template_paragraph._element)
    current_element = template_paragraph._element

    def next_paragraph():
        nonlocal current_element
        paragraph_element = deepcopy(template_element)
        current_element.addnext(paragraph_element)
        current_element = paragraph_element
        return Paragraph(paragraph_element, document._body)

    for index, node in enumerate(nodes):
        if isinstance(node, DocumentTitle):
            _add_paragraph(
                next_paragraph(),
                node.text,
                bold=True,
                size=HEADING_FONT_SIZE_PT,
                font_family=GOTHIC_FONT,
                indent_chars=1,
            )
            _add_blank_paragraph(next_paragraph())
            continue

        if isinstance(node, PartTitle):
            _add_paragraph(
                next_paragraph(),
                node.text,
                bold=True,
                size=HEADING_FONT_SIZE_PT,
                font_family=GOTHIC_FONT,
                indent_chars=1,
            )
            _add_blank_paragraph(next_paragraph())
            continue

        if not isinstance(node, Scene):
            continue

        _add_paragraph(
            next_paragraph(),
            f"○{node.text}",
            bold=True,
            size=HEADING_FONT_SIZE_PT,
            font_family=GOTHIC_FONT,
        )
        for block in node.blocks:
            if isinstance(block, Conversation):
                _add_conversation_paragraph(next_paragraph(), block)
            elif isinstance(block, StageDirection):
                _add_paragraph(
                    next_paragraph(),
                    block.text,
                    font_family=GOTHIC_FONT,
                    indent_chars=2,
                )
            elif isinstance(block, TextBlock):
                _add_paragraph(next_paragraph(), block.text)

        if any(isinstance(later, Scene) for later in nodes[index + 1 :]):
            _add_blank_paragraph(next_paragraph())

    template_paragraph._element.getparent().remove(template_paragraph._element)
    return document
